# -*- coding: utf-8 -*-
"""Chuyển Markdown -> .docx (Times New Roman 12, heading, bảng, list, bold/italic/code).

Cách dùng (PowerShell):
    python tools/md2docx.py docs/00_TongQuan_DuAn.md              # -> cùng tên .docx
    python tools/md2docx.py docs/*.md                             # nhiều file
    python tools/md2docx.py input.md output.docx                  # chỉ định tên ra
"""
import re
import sys
import glob
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(p, text):
    tokens = re.split(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'
        else:
            p.add_run(tok)


def md_to_doc(md, doc):
    lines = md.split('\n')
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run('\n'.join(code_buf))
                r.font.name = 'Consolas'; r.font.size = Pt(9)
                code_buf = []
            in_code = not in_code
            i += 1; continue
        if in_code:
            code_buf.append(lines[i]); i += 1; continue
        if not line.strip():
            i += 1; continue
        if re.match(r'^---+$', line.strip()):
            i += 1; continue
        if line.lstrip().startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                tbl_lines.append(lines[i].strip()); i += 1
            rows = []
            for tl in tbl_lines:
                if re.match(r'^\|[\s:|-]+\|?$', tl):
                    continue
                rows.append([c.strip() for c in tl.strip('|').split('|')])
            if rows:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncol)
                try:
                    t.style = 'Light Grid Accent 1'
                except Exception:
                    pass
                for ridx, r in enumerate(rows):
                    cells = t.add_row().cells
                    for ci in range(ncol):
                        para = cells[ci].paragraphs[0]
                        add_inline(para, r[ci] if ci < len(r) else '')
                        if ridx == 0:
                            for run in para.runs:
                                run.bold = True
            continue
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            h = doc.add_heading('', level=min(len(m.group(1)), 4))
            add_inline(h, m.group(2).strip())
            i += 1; continue
        if line.lstrip().startswith('>'):
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['Intense Quote']
            except Exception:
                pass
            add_inline(p, re.sub(r'^\s*>\s?', '', line))
            i += 1; continue
        if re.match(r'^\s*[-*]\s+', line):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, re.sub(r'^\s*[-*]\s+', '', line))
            i += 1; continue
        if re.match(r'^\s*\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            add_inline(p, re.sub(r'^\s*\d+\.\s+', '', line))
            i += 1; continue
        para_lines = [line]
        i += 1
        while (i < len(lines) and lines[i].strip()
               and not re.match(r'^(#{1,6}\s|>|\s*[-*]\s|\s*\d+\.\s|\||```)', lines[i].lstrip())
               and not re.match(r'^---+$', lines[i].strip())):
            para_lines.append(lines[i].rstrip()); i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, ' '.join(para_lines))


def build(md_path, out_path=None):
    if out_path is None:
        out_path = os.path.splitext(md_path)[0] + '.docx'
    md = open(md_path, encoding='utf-8').read()
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'; st.font.size = Pt(12)
    md_to_doc(md, doc)
    doc.save(out_path)
    print('Saved', out_path)


if __name__ == '__main__':
    args = sys.argv[1:]
    if not args:
        print(__doc__); sys.exit(1)
    if len(args) == 2 and args[1].lower().endswith('.docx'):
        build(args[0], args[1])
    else:
        for pattern in args:
            for f in (glob.glob(pattern) or [pattern]):
                if f.lower().endswith('.md'):
                    build(f)
